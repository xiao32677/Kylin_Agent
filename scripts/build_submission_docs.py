from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "dist" / "submission_docs"
PACK = ROOT / "dist" / "90016317介绍.zip"

BLUE = RGBColor(35, 100, 170)
DARK = RGBColor(23, 32, 42)
MUTED = RGBColor(90, 105, 122)
LIGHT = "F2F4F7"
PALE = "E8EEF5"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, widths: list[float] | None = None) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def configure_doc(doc: Document, title: str, subtitle: str) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 8),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(22)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(subtitle)
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(10.5)
    r.font.color.rgb = MUTED

    header = sec.header.paragraphs[0]
    header.text = "麒麟智维安全运维助手"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED

    footer = sec.footer.paragraphs[0]
    footer.text = "文档版本：V1.0"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_meta(doc: Document, doc_name: str) -> None:
    table = doc.add_table(rows=4, cols=2)
    style_table(table, [1.4, 5.1])
    rows = [
        ("文档名称", doc_name),
        ("系统名称", "麒麟智维安全运维助手"),
        ("运行环境", "麒麟 Linux / Python 3 / systemd"),
        ("适用对象", "系统管理员、运维人员、审计人员"),
    ]
    for row, (k, v) in zip(table.rows, rows):
        set_cell_text(row.cells[0], k, True)
        set_cell_text(row.cells[1], v)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths: list[float] = [1.75, 4.75]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    style_table(table, widths)
    for row, (k, v) in zip(table.rows, rows):
        set_cell_text(row.cells[0], k, True)
        set_cell_text(row.cells[1], v)


def add_matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    style_table(table, widths)
    for idx, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], h, True)
    for row_obj, row in zip(table.rows[1:], rows):
        for cell, text in zip(row_obj.cells, row):
            set_cell_text(cell, text)


def save(doc: Document, filename: str) -> Path:
    path = OUT / filename
    doc.save(path)
    return path


def requirement_doc() -> Path:
    doc = Document()
    configure_doc(doc, "软件功能需求分析文档", "面向安全智能运维 Agent 的功能、角色、安全和部署需求分析")
    add_meta(doc, "软件功能需求分析文档")
    doc.add_heading("1. 项目背景", level=1)
    doc.add_paragraph(
        "服务器运维场景中，管理员经常需要在资源异常、端口占用、日志告警、服务状态异常和磁盘空间不足等问题之间快速切换。"
        "传统命令行方式效率高但风险集中，误删、误重启、危险注入和越权操作都可能造成业务中断。系统需要在自然语言交互效率和安全可控之间取得平衡。"
    )
    doc.add_heading("2. 总体目标", level=1)
    add_bullets(doc, [
        "通过自然语言识别运维意图，并生成可审计的结构化执行计划。",
        "对危险命令、提示词注入、服务启停、权限变更、关键目录写入等行为进行风险过滤。",
        "基于最小权限原则运行执行代理，只允许白名单内的必要诊断命令。",
        "对高风险操作建立审批闭环，审批通过后才允许继续执行。",
        "将用户请求、安全检查、Tool 调用、审批、结果和知识沉淀写入审计链路。",
    ])
    doc.add_heading("3. 用户角色需求", level=1)
    add_matrix(doc, ["角色", "核心权限", "限制"], [
        ["admin", "对话运维、审批、审计、Tool 管理、知识库管理", "无业务功能限制"],
        ["operator", "对话运维、知识库、只读 Tool 信息", "不能审批，不能查看完整审计"],
        ["auditor", "审计追踪、知识库、只读信息", "不能执行运维对话和审批"],
    ], [1.15, 3.05, 2.3])
    doc.add_heading("4. 功能性需求", level=1)
    add_matrix(doc, ["编号", "需求项", "说明"], [
        ["FR-01", "自然语言对话", "识别资源、日志、端口、进程、服务、磁盘等运维意图。"],
        ["FR-02", "风险检查", "命令和文本输入均需经过规则库检查并输出风险等级。"],
        ["FR-03", "Tool 执行", "通过受控 Tool 采集系统状态，不直接拼接 shell 命令。"],
        ["FR-04", "审批中心", "高风险操作生成审批单，审批状态、过期时间和参数哈希必须校验。"],
        ["FR-05", "审计追踪", "按 trace_id 串联全流程，支持过滤和 CSV/JSONL 导出。"],
        ["FR-06", "知识库", "支持多层知识检索、引用和基于问答的知识沉淀。"],
        ["FR-07", "实时监控", "展示 CPU、内存、网络、磁盘、端口、审批和风险概况。"],
    ], [0.75, 1.6, 4.15])
    doc.add_heading("5. 非功能性需求", level=1)
    add_bullets(doc, [
        "安全性：默认拒绝危险操作；高风险变更必须审批；外部模型不可绕过护栏。",
        "可追溯性：关键请求必须具备 trace_id，并记录到 SQLite 与 JSONL 审计日志。",
        "可部署性：支持麒麟 Linux、systemd 服务、低权限账号和 sudoers 白名单。",
        "可维护性：安全规则支持配置化扩展，Tool 元数据包含名称、版本、权限、参数和健康状态。",
        "性能：命令风险检查 P95 目标小于 100ms，并支持不少于 20 个并发监控请求。",
    ])
    return save(doc, "01_软件功能需求分析文档.docx")


def design_doc() -> Path:
    doc = Document()
    configure_doc(doc, "软件功能设计文档", "系统架构、模块划分、数据流、接口和安全机制设计")
    add_meta(doc, "软件功能设计文档")
    doc.add_heading("1. 总体架构", level=1)
    doc.add_paragraph(
        "系统采用轻量 B/S 架构：前端控制台负责交互展示，后端 Python 服务负责任务编排、安全检查、Tool 调用、审批、审计和知识库。"
        "部署后由 systemd 托管，服务进程以 ops-agent 低权限账号运行。"
    )
    add_matrix(doc, ["层次", "模块", "职责"], [
        ["表现层", "frontend", "登录、对话运维、监控看板、知识库、审批中心、审计追踪、Tool 管理。"],
        ["编排层", "agent.py", "会话管理、意图处理、执行计划、审批创建、结果汇总、闭环报告。"],
        ["安全层", "guardrail.py", "文本与命令风险识别、路径策略、payload_hash。"],
        ["执行层", "tools.py", "资源、进程、端口、日志、服务、磁盘等受控 Tool。"],
        ["数据层", "storage.py", "SQLite 表结构、JSONL 审计、知识库和审批数据。"],
    ], [1.0, 1.55, 3.95])
    doc.add_heading("2. Agent 执行流程", level=1)
    add_numbers(doc, [
        "用户提交自然语言运维请求。",
        "系统生成 trace_id 并写入 user_message 审计事件。",
        "安全护栏对输入进行风险检查，危险请求直接拦截。",
        "本地规则或 DeepSeek 生成结构化意图和 Tool 计划。",
        "检索知识库并补充历史故障、运维规范和用户问答。",
        "只读 Tool 直接执行；高风险 Tool 进入审批中心。",
        "审批通过后校验状态、过期时间和 payload_hash，再继续执行。",
        "系统定位问题、生成建议、写入审计链路并沉淀知识。",
    ])
    doc.add_heading("3. Tool 设计", level=1)
    add_matrix(doc, ["Tool", "风险", "能力"], [
        ["get_system_overview", "readonly", "采集 OS、内核、架构、运行时间等信息。"],
        ["get_resource_usage", "readonly", "采集 CPU、内存、Swap 和负载指标。"],
        ["get_filesystem_usage", "readonly", "采集文件系统和挂载点使用率。"],
        ["find_large_files", "readonly", "扫描大文件并给出清理候选。"],
        ["list_ports", "readonly", "查询监听端口、进程和用户。"],
        ["query_journal", "readonly", "采集 systemd journal 告警和错误日志。"],
        ["find_deleted_open_files", "readonly", "检测已删除但仍被进程占用的文件。"],
        ["release_port_guarded", "high", "审批后释放被占用端口并验证结果。"],
        ["delete_file_guarded", "high", "审批后按路径策略删除文件。"],
        ["write_text_file_guarded", "medium", "审批后向受控目录写入文本文件。"],
    ], [2.0, 1.0, 3.5])
    doc.add_heading("4. 审批与审计设计", level=1)
    add_bullets(doc, [
        "审批表记录 approval_id、trace_id、风险等级、动作摘要、回滚方案、payload_hash、过期时间和状态。",
        "审批处理时必须满足 status=pending、未过期、payload_hash 未变化三个条件。",
        "审计事件覆盖 user_message、security_check、intent_classified、agent_plan、tool_invoked、approval_created、approval_decided、execution_result、knowledge_learned 等阶段。",
        "审计查询支持按用户、主机、风险等级、Tool、事件类型和时间范围过滤，并可导出 CSV/JSONL。",
    ])
    doc.add_heading("5. 部署与权限设计", level=1)
    add_kv_table(doc, [
        ("运行账号", "ops-agent 低权限系统账号"),
        ("服务目录", "/opt/a2-secops-agent"),
        ("配置目录", "/etc/a2-secops-agent"),
        ("数据目录", "/opt/a2-secops-agent/data"),
        ("服务托管", "systemd: a2-secops-agent.service"),
        ("sudoers 白名单", "journalctl、systemctl status、ss、lsof、dmesg"),
    ])
    return save(doc, "02_软件功能设计文档.docx")


def product_doc() -> Path:
    doc = Document()
    configure_doc(doc, "软件产品说明书", "系统功能、角色、操作入口和典型使用说明")
    add_meta(doc, "软件产品说明书")
    doc.add_heading("1. 产品概述", level=1)
    doc.add_paragraph(
        "麒麟智维安全运维助手面向服务器日常运维，提供对话式诊断、实时监控、安全审批、审计追踪、Tool 管理和知识库能力。"
        "系统重点解决“能快速运维”和“不能误操作”之间的矛盾。"
    )
    doc.add_heading("2. 登录与角色", level=1)
    add_matrix(doc, ["账号", "密码", "用途"], [
        ["admin", "a2admin123", "系统管理、审批、审计和配置展示。"],
        ["operator", "a2operator123", "日常运维对话和只读诊断。"],
        ["auditor", "a2auditor123", "审计查看和追溯分析。"],
    ], [1.4, 1.6, 3.5])
    doc.add_heading("3. 页面功能", level=1)
    add_matrix(doc, ["页面", "功能说明"], [
        ["对话运维", "输入自然语言问题，系统展示意图识别、安全过滤、Tool 调用、问题定位、知识引用和最终结论。"],
        ["知识库 RAG", "查看、检索、添加知识条目，系统也会根据问答自动沉淀知识。"],
        ["审计追踪", "按 trace_id 查看完整链路报告和事件时间线，支持导出。"],
        ["Tool 管理", "查看 Tool 注册信息、版本、权限、风险等级、启用状态和健康检查。"],
        ["审批中心", "处理高风险操作审批，审批通过后系统继续执行。"],
    ], [1.45, 5.05])
    doc.add_heading("4. 典型操作", level=1)
    add_numbers(doc, [
        "访问 http://127.0.0.1:8765 并使用 admin 登录。",
        "在对话运维输入“查看一下当前所有端口”。",
        "系统识别端口诊断意图并调用 list_ports Tool。",
        "在分析过程查看采集、分析、决策、执行、验证、知识和总结七步链路。",
        "如执行写文件、释放端口、删除文件等变更操作，进入审批中心处理。",
        "在审计追踪中按 trace_id 查看完整日志闭环。",
    ])
    doc.add_heading("5. 安全提示", level=1)
    add_bullets(doc, [
        "被安全护栏判定为 forbidden 的请求不会执行，也不会发送给外部模型。",
        "审批通过不等于无条件执行，系统仍会校验审批状态、过期时间和参数哈希。",
        "普通用户模式下系统级 Tool 会受当前用户权限限制，正式部署推荐使用 install_kylin.sh。",
        "低分辨率环境下可将浏览器缩放调整为 80% 或使用全屏模式。",
    ])
    return save(doc, "03_软件产品说明书.docx")


def function_test_doc() -> Path:
    doc = Document()
    configure_doc(doc, "软件功能测试报告", "基于当前实现的功能测试范围、用例和结果")
    add_meta(doc, "软件功能测试报告")
    doc.add_heading("1. 测试范围", level=1)
    add_bullets(doc, [
        "登录鉴权与 RBAC 权限控制。",
        "自然语言意图识别与安全风险过滤。",
        "资源、端口、CPU、服务和日志类诊断。",
        "审批创建、防重放和审批后执行。",
        "Tool 启用、禁用和健康检查。",
        "知识库检索、引用和问答沉淀。",
        "审计追踪、完整链路查询和导出。",
    ])
    doc.add_heading("2. 测试环境", level=1)
    add_kv_table(doc, [
        ("服务地址", "http://127.0.0.1:8765"),
        ("运行方式", "Python 后端服务 + Web 控制台"),
        ("测试脚本", "tests/smoke_test.py"),
        ("测试结论", "smoke test passed"),
    ])
    doc.add_heading("3. 主要测试用例", level=1)
    add_matrix(doc, ["编号", "用例", "预期结果", "结果"], [
        ["TC-01", "admin 登录", "返回 token 和用户信息", "通过"],
        ["TC-02", "获取 Tool 列表", "返回注册 Tool 元数据", "通过"],
        ["TC-03", "查看仪表盘", "返回资源和网络指标", "通过"],
        ["TC-04", "知识库检索 OOM Killer", "返回知识条目", "通过"],
        ["TC-05", "检查 systemctl restart sshd", "识别为 high/forbidden 并要求审批或阻断", "通过"],
        ["TC-06", "磁盘诊断对话", "返回 succeeded 或 pending_approval", "通过"],
        ["TC-07", "查看端口", "调用 list_ports 并返回端口数据", "通过"],
        ["TC-08", "危险注入 rm/chmod", "security.allowed=false 并写入审计", "通过"],
        ["TC-09", "审批重复 approve", "第二次返回 SEC_APPROVAL_INVALID", "通过"],
        ["TC-10", "Tool 禁用/启用", "状态可变更且健康检查可用", "通过"],
        ["TC-11", "operator 权限", "可对话但受审批/审计权限限制", "通过"],
    ], [0.65, 1.95, 3.05, 0.85])
    doc.add_heading("4. 测试结论", level=1)
    doc.add_paragraph(
        "当前版本已通过核心功能冒烟测试，覆盖登录、RBAC、Tool 注册、资源监控、知识库、安全检查、审批防重放、审计查询和典型对话运维流程。"
    )
    return save(doc, "04_软件功能测试报告.docx")


def performance_doc() -> Path:
    doc = Document()
    configure_doc(doc, "软件性能测试报告", "核心指标、测试方法、实测结果和结论")
    add_meta(doc, "软件性能测试报告")
    doc.add_heading("1. 测试目标", level=1)
    add_bullets(doc, [
        "验证命令风险检查接口在高频调用下的响应能力。",
        "验证实时监控接口在 20 并发下的可用性。",
        "采集运行时资源样本，为部署容量评估提供参考。",
    ])
    doc.add_heading("2. 测试方法", level=1)
    add_kv_table(doc, [
        ("测试脚本", "tests/performance_test.py"),
        ("服务地址", "http://127.0.0.1:8765"),
        ("风险检查样本", "rm -rf /、chmod -R 777 /etc、systemctl restart sshd 循环 60 次"),
        ("并发场景", "20 个并发请求访问 /api/v1/dashboard"),
        ("统计指标", "平均耗时、P50、P95、最大耗时"),
    ])
    doc.add_heading("3. 实测结果", level=1)
    add_matrix(doc, ["指标", "次数", "平均", "P50", "P95", "最大值", "目标"], [
        ["guardrail_check", "60", "18.20ms", "17.12ms", "31.57ms", "43.89ms", "P95 < 100ms"],
        ["dashboard_20_concurrent", "20", "1543.40ms", "1858.82ms", "2333.09ms", "2413.17ms", "P95 < 5000ms"],
    ], [1.65, 0.6, 0.85, 0.85, 0.85, 0.9, 1.3])
    doc.add_heading("4. 资源样本", level=1)
    add_kv_table(doc, [
        ("CPU 使用率样本", "40.8%"),
        ("内存使用率样本", "78.1%"),
        ("并发能力", "20 并发监控请求完成且无错误"),
    ])
    doc.add_heading("5. 结论", level=1)
    doc.add_paragraph(
        "风险检查接口 P95 为 31.57ms，满足小于 100ms 的目标；20 并发监控接口 P95 为 2333.09ms，满足小于 5000ms 的目标。"
        "当前性能能够支撑演示和小规模运维管理场景。"
    )
    return save(doc, "05_软件性能测试报告.docx")


def install_doc() -> Path:
    doc = Document()
    configure_doc(doc, "软件安装包及部署文档", "发布包内容、麒麟虚拟机安装、启动、配置和排障步骤")
    add_meta(doc, "软件安装包及部署文档")
    doc.add_heading("1. 发布包说明", level=1)
    add_kv_table(doc, [
        ("安装包名称", "90016317作品.zip"),
        ("源码包名称", "90016317源码.zip"),
        ("访问地址", "http://127.0.0.1:8765"),
        ("启动方式", "sudo bash deploy/install_kylin.sh"),
    ])
    doc.add_heading("2. 安装步骤", level=1)
    add_numbers(doc, [
        "打开麒麟虚拟机终端，进入发布包所在目录，例如 cd ~/Downloads。",
        "执行 unzip a2-secops-agent-release-*.zip 解压发布包。",
        "进入目录：cd a2-secops-agent。",
        "执行 sudo bash deploy/install_kylin.sh 安装并启动服务。",
        "执行 systemctl status a2-secops-agent --no-pager 检查服务状态。",
        "在浏览器访问 http://127.0.0.1:8765。",
    ])
    doc.add_heading("3. 安装脚本动作", level=1)
    add_bullets(doc, [
        "创建 ops-agent 低权限运行账号和用户组。",
        "将程序安装到 /opt/a2-secops-agent。",
        "创建 /etc/a2-secops-agent 配置目录。",
        "安装 security_rules.json 安全规则库。",
        "安装 /etc/sudoers.d/a2-secops-agent 白名单。",
        "安装 /etc/systemd/system/a2-secops-agent.service 并启动服务。",
    ])
    doc.add_heading("4. 常用命令", level=1)
    add_matrix(doc, ["场景", "命令"], [
        ["启动服务", "sudo systemctl start a2-secops-agent"],
        ["重启服务", "sudo systemctl restart a2-secops-agent"],
        ["停止服务", "sudo systemctl stop a2-secops-agent"],
        ["设置自启", "sudo systemctl enable a2-secops-agent"],
        ["查看日志", "journalctl -u a2-secops-agent -n 100 --no-pager"],
        ["健康检查", "curl http://127.0.0.1:8765/api/v1/health"],
    ], [1.4, 5.1])
    doc.add_heading("5. DeepSeek 配置", level=1)
    doc.add_paragraph("编辑 /etc/a2-secops-agent/a2-secops-agent.env，写入以下配置后重启服务：")
    add_kv_table(doc, [
        ("DEEPSEEK_API_KEY", "your_api_key"),
        ("DEEPSEEK_BASE_URL", "https://api.deepseek.com"),
        ("DEEPSEEK_MODEL", "deepseek-v4-flash"),
    ])
    doc.add_heading("6. 普通用户启动", level=1)
    doc.add_paragraph(
        "如果当前环境不能使用 sudo 或 systemd，可以执行 bash deploy/run_demo_user.sh。普通用户方式保留 Web 控制台、鉴权、RBAC、风险过滤、审批、审计、知识库和只读诊断能力。"
    )
    return save(doc, "06_软件安装包及部署文档.docx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    for old in OUT.glob("*.docx"):
        old.unlink()
    paths = [
        requirement_doc(),
        design_doc(),
        product_doc(),
        function_test_doc(),
        performance_doc(),
        install_doc(),
    ]
    if PACK.exists():
        PACK.unlink()
    with ZipFile(PACK, "w", ZIP_DEFLATED) as zf:
        for path in paths:
            zf.write(path, path.name)
    print(PACK)
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()

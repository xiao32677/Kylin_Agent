from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "dist" / "submission_docs"


def main() -> None:
    for path in sorted(DOCS.glob("*.docx")):
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        print(path.name)
        print("  paragraphs:", len(paragraphs))
        print("  tables:", len(doc.tables))
        print("  title:", paragraphs[0] if paragraphs else "-")
        if len(paragraphs) < 10:
            raise SystemExit(f"{path.name} content too short")
        if not doc.tables:
            raise SystemExit(f"{path.name} missing tables")


if __name__ == "__main__":
    main()
